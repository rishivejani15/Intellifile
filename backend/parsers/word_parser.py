import zipfile
from docx import Document

def extract_word_structure(file_path):
    """
    Extracts paragraphs, headings, and tables from a .docx file.
    """
    has_macros = False
    try:
        with zipfile.ZipFile(file_path, 'r') as z:
            if 'word/vbaProject.bin' in z.namelist():
                has_macros = True
    except Exception:
        pass

    try:
        doc = Document(file_path)

        paragraphs = []
        headings = []
        tables = []

        # 1. Extract Normal Paragraphs
        for i, para in enumerate(doc.paragraphs):
            # Extract raw text first - THIS IS SACRED
            raw_text = para.text.strip()
            
            # Look for special objects (Images, Breaks)
            has_image = any('w:drawing' in r.element.xml or 'w:pict' in r.element.xml for r in para.runs)
            has_break = any('w:br' in r.element.xml or 'w:lastRenderedPageBreak' in r.element.xml for r in para.runs)
            
            final_text = raw_text
            
            # If there's no text but there is a graphic, tag it
            if not raw_text:
                if has_image:
                    final_text = "[IMAGE / GRAPHIC]"
                elif has_break:
                    final_text = "[PAGE BREAK / SECTION]"
                else:
                    # Truly empty line with no objects - skip to avoid clutter
                    continue
            else:
                # If there IS text AND an image in the same block, keep the text!
                # (Optional: append a small tag if you want to know an image was next to the text)
                if has_image:
                    final_text = f"{raw_text} [Graphic Attached]"

            # Track headings for the forensic summary
            if para.style.name.startswith("Heading"):
                headings.append(final_text)
            
            paragraphs.append(final_text)
        
        # 1.5. POST-EXTRACTION SANITIZATION (The 'Anti-Ghosting' Filter)
        # If the document starts with empty lines (often caused by Word restoration/formatting), 
        # we strip them so that the FIRST LINE of real text is always index 0.
        while paragraphs and (not paragraphs[0] or paragraphs[0] in ["[IMAGE / GRAPHIC]", "[PAGE BREAK / SECTION]"]):
            if not paragraphs[0].strip():
                paragraphs.pop(0)
            else:
                break

        # 2. Extract Table Content (Crucial for bordered/tabular docs)
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        # Add table text to main paragraphs so it's tracked in the diff!
                        paragraphs.append(f"[TABLE CELL]: {cell_text}")
                        row_cells.append(cell_text)
                table_data.append(row_cells)
            tables.append(table_data)

        return {
            "paragraphs": paragraphs,
            "headings": headings,
            "tables": tables,
            "has_macros": has_macros
        }
    except Exception as e:
        print(f"[WordParser] Error: {str(e)}")
        return {
            "paragraphs": [],
            "headings": [],
            "tables": [],
            "has_macros": has_macros
        }
