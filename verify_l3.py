import os
import sys
from docx import Document
from openpyxl import Workbook

# Add backend to sys.path
sys.path.append(os.path.join(os.getcwd(), 'backend'))

from core.versioning.version_engine import VersionEngine

def create_dummy_docx(path, paragraphs, headings):
    doc = Document()
    for h in headings:
        doc.add_heading(h, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(path)

def create_dummy_xlsx(path, sheet_name, cell_data):
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name
    for coord, value in cell_data.items():
        ws[coord] = value
    wb.save(path)

def test_word_versioning():
    print("\n--- Testing Word Versioning ---")
    engine = VersionEngine()
    
    old_path = "test_old.docx"
    new_path = "test_new.docx"
    
    create_dummy_docx(old_path, ["Para 1", "Para 2"], ["Heading 1"])
    create_dummy_docx(new_path, ["Para 1", "Para 3"], ["Heading 1", "Heading 2"])
    
    result = engine.process_version("test.docx", old_path, new_path)
    print(f"Summary: {result['summary']}")
    print(f"Risk: {result['risk_level']}")
    print(f"Diff: {result['diff']}")
    
    os.remove(old_path)
    os.remove(new_path)

def test_excel_versioning():
    print("\n--- Testing Excel Versioning ---")
    engine = VersionEngine()
    
    old_path = "test_old.xlsx"
    new_path = "test_new.xlsx"
    
    create_dummy_xlsx(old_path, "Sheet1", {"A1": "Val1", "B1": "Val2"})
    create_dummy_xlsx(new_path, "Sheet1", {"A1": "Val1", "B1": "Val3", "C1": "Val4"})
    
    result = engine.process_version("test.xlsx", old_path, new_path)
    print(f"Summary: {result['summary']}")
    print(f"Risk: {result['risk_level']}")
    print(f"Diff: {result['diff']}")
    
    os.remove(old_path)
    os.remove(new_path)

if __name__ == "__main__":
    try:
        test_word_versioning()
        test_excel_versioning()
        print("\nVerification scripts finished successfully.")
    except Exception as e:
        print(f"\nVerification failed: {str(e)}")
        import traceback
        traceback.print_exc()
